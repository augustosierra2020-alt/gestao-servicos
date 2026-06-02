import io
import re
import pandas as pd
import streamlit as st
from openpyxl.styles import PatternFill
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuração da página do Streamlit
st.set_page_config(
    page_title="Gestão de Serviços & OS - Hyper Tork", page_icon="📊", layout="wide"
)

st.title("📊 Gestão de Serviços & Emissão de OS")
st.write(
    "Filtragem, cálculo de valores, remoção de duplicadas por Matrícula e preenchimento automático do modelo Word da Hyper Tork."
)

# Criando abas para organizar o fluxo do sistema
aba1, aba2 = st.tabs(["📋 Processamento da Planilha", "📄 Gerar Ordem de Serviço"])

if "df_filtrado" not in st.session_state:
    st.session_state.df_filtrado = None

# --- FUNÇÃO DE CÁLCULO DE VALOR ---
def calcular_valor_inicial(linha):
    descricao = str(linha.get("Nome arquivo", "")).upper().strip()
    veiculo = str(linha.get("Fabricante", "")).upper().strip()

    descricao = re.sub(r"\s+", " ", descricao)
    veiculo = re.sub(r"\s+", " ", veiculo)

    termos_stg2 = ["STG2", "STG 2", "STAG2", "STAG 2"]
    termos_mod_off = ["MOD", "OFF"]

    fabricantes_especiais = [
        "NEW HOLLAND", "VALTRA", "CASE IH", "CASE", "MASSEY FERGUSSON", 
        "MASSEY", "CLAAS", "JHON DEERE", "JOHN DEERE", "DEERE", 
        "FENDT", "JACTO", "DOPPSTADT", "JAN", "VOLVO CONSTRUCTION EQUIPMENT", 
        "VOLVO CONSTRUCTION", "VOLVO CE"
    ]

    eh_especial = any(fab in veiculo for fab in fabricantes_especiais)
    if "VOLVO TRUCK" in veiculo:
        eh_especial = False

    if any(termo in descricao for termo in termos_stg2):
        return 1400 if eh_especial else 650
    elif any(termo in descricao for termo in termos_mod_off):
        return 700 if eh_especial else 350

    return None

# --- FUNÇÃO PARA FILTRAR APENAS OS TERMOS SOLICITADOS NA DESCRIÇÃO ---
def limpar_descricao_os(desc_original):
    desc_upper = str(desc_original).upper().strip()
    if "STAG 2" in desc_upper or "STAG2" in desc_upper:
        return "STAG 2"
    elif "STG 2" in desc_upper or "STG2" in desc_upper:
        return "STG 2"
    elif "MOD" in desc_upper:
        return "MOD"
    elif "OFF" in desc_upper:
        return "OFF"
    return desc_original

# --- FUNÇÃO QUE EDITA DIRETAMENTE O SEU ARQUIVO ORIGINAL ---
def modificar_modelo_docx(modelo_bytes, flash_point, cliente_nome, cidade, contato, linhas_tabela, total_valor):
    doc = Document(io.BytesIO(modelo_bytes))
    
    # 1. Preencher os dados do Cabeçalho estritamente na caixa/célula da frente
    for t in doc.tables:
        for row in t.rows:
            if len(row.cells) >= 2:
                texto_celula_1 = row.cells[0].text.upper().strip()
                
                if "CLIENTE:" in texto_celula_1:
                    row.cells[1].text = f"{cliente_nome} - {flash_point}"
                    for p in row.cells[1].paragraphs:
                        for run in p.runs: run.font.name = 'Arial'; run.font.size = Pt(11)
                        
                elif "CIDADE:" in texto_celula_1:
                    row.cells[1].text = cidade
                    for p in row.cells[1].paragraphs:
                        for run in p.runs: run.font.name = 'Arial'; run.font.size = Pt(11)
                        
                elif "CONTATO:" in texto_celula_1:
                    row.cells[1].text = contato
                    for p in row.cells[1].paragraphs:
                        for run in p.runs: run.font.name = 'Arial'; run.font.size = Pt(11)

    # 2. Preencher a tabela de serviços original do seu documento
    linhas_validas = [l for l in linhas_tabela if l.get("Valor") is not None and str(l.get("Valor")).strip() != "" and str(l.get("Valor")).lower() != "nan"]
    
    tabela_servicos = None
    for t in doc.tables:
        if len(t.rows) > 0 and "Nº MAPA" in t.rows[0].cells[0].text.upper():
            tabela_servicos = t
            break
            
    if tabela_servicos:
        for i, linha in enumerate(linhas_validas):
            idx_linha_destino = i + 1
            
            if idx_linha_destino >= len(tabela_servicos.rows):
                row_cells = tabela_servicos.add_row().cells
            else:
                row_cells = tabela_servicos.rows[idx_linha_destino].cells
                
            dados_linha = [
                str(linha.get("Nº Mapa", "")),
                str(linha.get("Data", "")),
                str(linha.get("Veículo", "")),
                str(linha.get("Placa", "")),
                limpar_descricao_os(linha.get("Descrição", "")),
                f"R$ {linha.get('Valor', '')}"
            ]
            
            for idx_col, valor_celula in enumerate(dados_linha):
                if idx_col < len(row_cells):
                    row_cells[idx_col].text = valor_celula
                    for p_cell in row_cells[idx_col].paragraphs:
                        if idx_col in [0, 1, 3, 5]:
                            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p_cell.runs:
                            r.font.name = 'Arial'
                            r.font.size = Pt(10)
                            
        # EXCLUSÃO DAS LINHAS SOBRESSALENTES VAZIAS
        linha_inicio_remocao = len(linhas_validas) + 1
        while len(tabela_servicos.rows) > linha_inicio_remocao:
            linha_para_apagar = tabela_servicos.rows[linha_inicio_remocao]
            tabela_servicos._tbl.remove(linha_para_apagar._tr)

    # 3. INTERAÇÃO PARA O CAMPO DO TOTAL
    if pd.isna(total_valor) or str(total_valor).lower() == "nan":
        total_valor = 0.0

    valor_formatado_texto = f"{float(total_valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    
    for t in doc.tables:
        for row in t.rows:
            contem_total = any("TOTAL" in cell.text.upper() for cell in row.cells)
            if contem_total:
                for cell in row.cells:
                    if "R$" in cell.text or "NAN" in cell.text.upper() or cell.text.strip() == "":
                        cell.text = f"R$ {valor_formatado_texto}"
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in p.runs:
                                run.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(12)
                                run.font.color.rgb = RGBColor(234, 88, 12)

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target


# ==============================================================================
# --- ABA 1: PROCESSAMENTO DA PLANILHA ---
# ==============================================================================
with aba1:
    arquivo_carregado = st.file_uploader(
        "Arraste ou selecione a planilha FPF_List para iniciar:",
        type=["xlsx", "xls", "csv"],
        key="uploader_planilha"
    )

    if arquivo_carregado is not None:
        try:
            conteudo = arquivo_carregado.read()
            try:
                excel_file = pd.ExcelFile(io.BytesIO(conteudo))
                abas = excel_file.sheet_names
                if len(abas) > 1:
                    aba_selecionada = st.selectbox("Selecione a aba com os dados:", abas, key="selecao_abas_app")
                else:
                    aba_selecionada = abas[0]
                df = pd.read_excel(io.BytesIO(conteudo), sheet_name=aba_selecionada)
            except Exception:
                try:
                    df = pd.read_csv(io.BytesIO(conteudo), sep=";", encoding="utf-8")
                    if df.shape[1] <= 1:
                        df = pd.read_csv(io.BytesIO(conteudo), sep=",", encoding="utf-8")
                except Exception:
                    df = pd.read_csv(io.BytesIO(conteudo), sep=";", encoding="iso-8859-1")

            if df is None or df.empty or len(df.columns) == 0:
                st.error("Erro: Não foi possível processar a estrutura de dados deste arquivo.")
            else:
                df.columns = df.columns.str.strip()

                if "T" in df.columns:
                    df["T"] = df["T"].astype(str).str.strip()
                    df_filtrado = df[df["T"] == "MOD"].copy()
                else:
                    st.warning("Aviso: A coluna 'T' não foi encontrada.")
                    df_filtrado = df.copy()

                # Adicionado "Cliente" explicitamente na captura inicial de colunas
                colunas_originais = ["Arquivo ID", "Fabricante", "Matrícula", "FlashPoint", "Cliente", "Nome arquivo", "Dada"]
                colunas_existentes = [col for col in colunas_originais if col in df_filtrado.columns]
                df_filtrado = df_filtrado[colunas_existentes].copy()

                df_filtrado["Valor"] = df_filtrado.apply(calcular_valor_inicial, axis=1)

                dicionario_renomear = {
                    "Arquivo ID": "Nº Mapa",
                    "Fabricante": "Veículo",
                    "Matrícula": "Placa",
                    "Nome arquivo": "Descrição",
                    "Dada": "Data",
                    "FlashPoint": "Flash Point",
                }
                df_filtrado = df_filtrado.rename(columns=dicionario_renomear)

                # --- ALTERAÇÃO SOLICITADA: Incluído 'Cliente' exatamente entre 'Flash Point' e 'Descrição' ---
                ordem_solicitada = ["Nº Mapa", "Data", "Veículo", "Placa", "Flash Point", "Cliente", "Descrição", "Valor"]
                colunas_finais = [col for col in ordem_solicitada if col in df_filtrado.columns]
                df_filtrado = df_filtrado[colunas_finais].copy()

                # Ajuste e formatação da data para o padrão Brasil
                if "Data" in df_filtrado.columns:
                    df_filtrado["Data"] = pd.to_datetime(df_filtrado["Data"], errors='coerce')
                    df_filtrado["Data"] = df_filtrado["Data"].dt.strftime('%d/%m/%Y').fillna("")

                # Força a coluna Flash Point a virar Texto/String pura
                if "Flash Point" in df_filtrado.columns:
                    df_filtrado["Flash Point"] = df_filtrado["Flash Point"].astype(str).str.strip()
                    df_filtrado = df_filtrado.sort_values(by=["Flash Point", "Nº Mapa"] if "Nº Mapa" in df_filtrado.columns else ["Flash Point"], ascending=True)

                st.session_state.df_filtrado = df_filtrado.copy()

                if not df_filtrado.empty and "Flash Point" in df_filtrado.columns:
                    lista_linhas = []
                    linhas_amarelas = []
                    linhas_laranjas = []
                    contador_linha_excel = 2

                    for fp, bloco in df_filtrado.groupby("Flash Point", sort=False):
                        placas_vistas = set()

                        for idx, linha in bloco.iterrows():
                            linha_dict = linha.to_dict()
                            placa_atual = str(linha.get("Placa", "")).strip()

                            if placa_atual in placas_vistas and placa_atual != "":
                                linha_dict["Valor"] = None
                                linhas_amarelas.append(contador_linha_excel)
                            else:
                                if placa_atual != "":
                                    placas_vistas.add(placa_atual)

                            lista_linhas.append(linha_dict)
                            contador_linha_excel += 1

                        df_bloco_temp = pd.DataFrame(lista_linhas[-len(bloco) :])
                        soma_bloco = pd.to_numeric(df_bloco_temp["Valor"], errors="coerce").sum()

                        linha_total = {col: "" for col in colunas_finais}
                        linha_total["Flash Point"] = fp
                        # Se a coluna 'Cliente' existir no mapeamento final, mantém o nome do cliente na linha do total
                        if "Cliente" in linha_total:
                            linha_total["Cliente"] = str(bloco.iloc[0].get("Cliente", ""))
                        linha_total["Descrição"] = "VALOR TOTAL:"
                        linha_total["Valor"] = float(soma_bloco) if soma_bloco > 0 else ""

                        lista_linhas.append(linha_total)
                        linhas_laranjas.append(contador_linha_excel)
                        contador_linha_excel += 1

                        linha_espacamento = {col: "" for col in colunas_finais}
                        lista_linhas.append(linha_espacamento)
                        contador_linha_excel += 1

                    if lista_linhas:
                        lista_linhas.pop()

                    df_excel_final = pd.DataFrame(lista_linhas, columns=colunas_finais)

                    st.subheader("📋 Visualização Prévia dos Dados Processados")
                    st.dataframe(df_filtrado)

                    buffer = io.BytesIO()
                    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
                        df_excel_final.to_excel(writer, index=False, sheet_name="FPF Realizados")

                        workbook = writer.book
                        worksheet = writer.sheets["FPF Realizados"]

                        amarelo_claro = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")
                        laranja_claro = PatternFill(start_color="FFE6CC", end_color="FFE6CC", fill_type="solid")

                        for num_linha in linhas_amarelas:
                            for col_idx in range(1, len(colunas_finais) + 1):
                                worksheet.cell(row=num_linha, column=col_idx).fill = amarelo_claro

                        for num_linha in linhas_laranjas:
                            for col_idx in range(1, len(colunas_finais) + 1):
                                worksheet.cell(row=num_linha, column=col_idx).fill = laranja_claro

                    st.success("Planilha processada com sucesso na memória! Vá para a aba ao lado para gerar Ordens de Serviço.")
                    st.download_button(
                        label="📥 Baixar Planilha Processada (Excel)",
                        data=buffer.getvalue(),
                        file_name="FPF_Relatorio_Final.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )
        except Exception as e:
            st.error(f"Erro crítico no processamento: {e}")

# ==============================================================================
# --- ABA 2: PREENCHIMENTO E CONSULTA DA ORDEM DE SERVIÇO ---
# ==============================================================================
with aba2:
    st.subheader("📄 Emissor de Ordem de Serviço com Base no Modelo Original")
    
    modelo_word_carregado = st.file_uploader(
        "Selecione o seu arquivo original 'MODELO - HYPER TORK PERFORMANCE.docx':",
        type=["docx"],
        key="uploader_modelo_word"
    )
    
    if st.session_state.df_filtrado is None or st.session_state.df_filtrado.empty:
        st.info("Aguardando o upload e processamento da planilha FPF_List na primeira aba para liberar o emissor.")
    elif modelo_word_carregado is None:
        st.warning("Por favor, anexe o arquivo original do seu modelo acima para habilitar o preenchimento automático.")
    else:
        df_base_os = st.session_state.df_filtrado
        bytes_modelo = modelo_word_carregado.read()
        
        lista_fp_unicos = sorted(list(set(str(val).strip() for val in df_base_os["Flash Point"].unique() if pd.notna(val))))
        
        fp_selecionado = st.selectbox("Selecione o Flash Point para gerar a OS correspondente:", lista_fp_unicos)
        
        dados_bloco = df_base_os[df_base_os["Flash Point"] == fp_selecionado]
        cliente_sugerido = str(dados_bloco.iloc[0].get("Cliente", "Cliente Não Identificado"))
        
        col1, col2 = st.columns(2)
        with col1:
            nome_cliente_input = st.text_input("Cliente (Preenchido Automaticamente):", value=cliente_sugerido)
            cidade_input = st.text_input("Cidade (Adicionar a critério do usuário):", placeholder="Ex: Cascavel - PR")
        with col2:
            flash_point_confirmacao = st.text_input("Flash Point Relacionado:", value=fp_selecionado, disabled=True)
            contato_input = st.text_input("Contato (Adicionar a critério do usuário):", placeholder="Ex: (45) 99999-9999")
            
        st.write("### Serviços com Valores Definidos que farão parte desta OS (Linhas vazias serão excluídas do Word):")
        
        linhas_os_finais = []
        placas_vistas_os = set()
        soma_total_os = 0
        
        for idx, row in dados_bloco.iterrows():
            row_dict = row.to_dict()
            placa = str(row_dict.get("Placa", "")).strip()
            
            if placa in placas_vistas_os and placa != "":
                row_dict["Valor"] = None
            else:
                if placa != "":
                    placas_vistas_os.add(placa)
                if row_dict["Valor"] is not None and str(row_dict["Valor"]).lower() != "nan" and not pd.isna(row_dict["Valor"]):
                    soma_total_os += float(row_dict["Valor"])
            
            linhas_os_finais.append(row_dict)
            
        df_preview_os = pd.DataFrame(linhas_os_finais)
        df_preview_os = df_preview_os[df_preview_os["Valor"].notna()].copy()
        df_preview_os["Descrição"] = df_preview_os["Descrição"].apply(limpar_descricao_os)
        
        # Exibição do preview mantendo o controle das colunas que vão para a tela
        colunas_preview_os = ["Nº Mapa", "Data", "Veículo", "Placa", "Descrição", "Valor"]
        colunas_preview_existentes = [c for c in colunas_preview_os if c in df_preview_os.columns]
        st.dataframe(df_preview_os[colunas_preview_existentes])
        
        st.metric(label="Valor Total Consolidado da OS (Apenas linhas válidas)", value=f"R$ {soma_total_os:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
        
        if st.button("🚀 Preencher e Gerar Ordem de Serviço"):
            arquivo_word_final = modificar_modelo_docx(
                modelo_bytes=bytes_modelo,
                flash_point=fp_selecionado,
                cliente_nome=nome_cliente_input,
                cidade=cidade_input,
                contato=contato_input,
                linhas_tabela=linhas_os_finais,
                total_valor=soma_total_os
            )
            
            st.success(f"Ordem de Serviço para o Flash Point {fp_selecionado} gerada com sucesso!")
            
            st.download_button(
                label="📥 Baixar Ordem de Serviço Pronta (.docx)",
                data=arquivo_word_final.getvalue(),
                file_name=f"OS_Hyper_Tork_{fp_selecionado}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )